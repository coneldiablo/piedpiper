"""
ThreatInquisitor/analyzer/static_analysis.py

Действительно улучшенный модуль статического анализа:
- PE (Windows .exe/.dll) через pefile + capstone-дизассемблирование
- ELF (Linux бинарь) через pyelftools + capstone-дизассемблирование
- PDF через PyPDF2
- DOCX через python-docx
- Старые DOC (OLE) через oletools (VBA_Parser)
- Скрипты (PowerShell, VBS, JS) - базовый анализ на подозрительные конструкции.
- Хэши MD5/SHA256
- Детект EICAR
- Интеграция с Yara для сканирования по правилам.
- Новая логика: "enhanced checks"
   * Проверяем секции на признаки упаковки (UPX, высокая энтропия, мало импортов, сигнатуры)
   * Ищем подозрительные API (CreateRemoteThread, VirtualAllocEx и др.)
   * Сканируем строки на IP/URL
+ - Улучшено обнаружение упаковщиков (энтропия, импорты, базовые сигнатуры).
"""

import os
import re
import math # Для энтропии
import base64
import zipfile
import collections # Для энтропии
import logging
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple # Добавлено для Yara
from xml.etree import ElementTree as ET

import hashlib
try:
    import pefile
except ImportError:
    pefile = None
try:
    import networkx as nx
except ImportError:
    nx = None
try:
    from elftools.elf.elffile import ELFFile
except ImportError:
    ELFFile = None
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    import docx
except ImportError:
    docx = None
try:
    from oletools.olevba import VBA_Parser, FileOpenError
except ImportError:
    VBA_Parser = None

# Capstone
try:
    import capstone
except ImportError:
    capstone = None

# Yara
logger = logging.getLogger("static_analysis")
logger.setLevel(logging.DEBUG)
try:
    import yara
except ImportError:
    yara = None
    logger.warning("Модуль yara не найден. Yara-сканирование будет недоступно.")

# Regex для IoC в скриптах (дублируем из ioc_extractor для простоты)
IP_REGEX_SCRIPT = re.compile(r"\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b")
URL_REGEX_SCRIPT = re.compile(r"(?:https?|ftp)://[\w\-._~:/?#[\]@!$&'()*+,;=%]+", re.IGNORECASE)
# === Enhanced heuristics constants ===
SUSPICIOUS_IMPORT_NAMES = {'VirtualAllocEx', 'CreateRemoteThread', 'WriteProcessMemory', 'VirtualProtect'}
SUSPICIOUS_DLL_NAMES = {'kernel32.dll', 'advapi32.dll', 'ntdll.dll'}
PACKER_SECTION_NAMES = {'.upx', 'UPX0', 'UPX1'}
SUSPICIOUS_KEYWORDS = ['psexec', 'mimikatz', 'powershell', 'reverse shell']
HIGH_ENTROPY_THRESHOLD = 7.2
PACKER_SECTION_NAMES_LOWER = {name.lower() for name in PACKER_SECTION_NAMES}

# Настройки расширенного анализа строк
STRING_MIN_LENGTH = 4
STRING_EXTRACTION_LIMIT = 1500
STRING_SAMPLE_SIZE = 25
BASE64_ALPHABET = set("ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=")
EMAIL_REGEX = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")


###############################################################################
# 1. Хелпер: подсчёт MD5/SHA256 и Энтропии
###############################################################################

def get_file_hashes(filepath: str) -> dict:
    md5h = hashlib.md5()
    sha256h = hashlib.sha256()
    with open(filepath, "rb") as f:
        while True:
            chunk = f.read(8192)
            if not chunk:
                break
            md5h.update(chunk)
            sha256h.update(chunk)
    return {
        "md5": md5h.hexdigest(),
        "sha256": sha256h.hexdigest()
    }

def calculate_entropy(data: bytes) -> float:
    """ Рассчитывает энтропию Шеннона для байтовой строки. """
    if not data:
        return 0.0
    entropy = 0
    data_len = len(data)
    byte_counts = collections.Counter(data)

    for count in byte_counts.values():
        # p_x = вероятность байта x
        p_x = count / data_len
        entropy += - p_x * math.log2(p_x)

    return entropy

# Примечание: ранее здесь была функция get_file_entropy(),
# которая не использовалась нигде в проекте. Удалена как лишняя.


def _deduplicate_preserve_order(items: List[str]) -> List[str]:
    """Удаляет дубликаты, сохраняя порядок появления."""
    seen = set()
    unique_items = []
    for item in items:
        if item not in seen:
            seen.add(item)
            unique_items.append(item)
    return unique_items


def _extract_ascii_strings(data: bytes, min_length: int) -> List[str]:
    current = []
    result: List[str] = []
    for byte in data:
        if 32 <= byte <= 126:
            current.append(chr(byte))
        else:
            if len(current) >= min_length:
                result.append(''.join(current))
            current = []
    if len(current) >= min_length:
        result.append(''.join(current))
    return result


def _extract_utf16le_strings(data: bytes, min_length: int) -> List[str]:
    current = []
    result: List[str] = []
    data_len = len(data) - (len(data) % 2)
    for idx in range(0, data_len, 2):
        char = data[idx]
        null_byte = data[idx + 1]
        if null_byte == 0 and 32 <= char <= 126:
            current.append(chr(char))
        else:
            if len(current) >= min_length:
                result.append(''.join(current))
            current = []
    if len(current) >= min_length:
        result.append(''.join(current))
    return result


def extract_strings(filepath: str,
                    min_length: int = STRING_MIN_LENGTH,
                    limit: int = STRING_EXTRACTION_LIMIT) -> Dict[str, Any]:
    """Извлекает строки (ASCII и UTF-16LE) из файла для дальнейшего анализа."""
    if not os.path.isfile(filepath):
        return {"strings": [], "truncated": False}

    file_size = os.path.getsize(filepath)
    max_bytes = 8 * 1024 * 1024  # Читаем не более 8 МБ для производительности
    truncated = file_size > max_bytes

    try:
        with open(filepath, "rb") as f:
            data = f.read(max_bytes)
    except Exception as exc:
        logger.error(f"Не удалось извлечь строки из {filepath}: {exc}")
        return {"strings": [], "truncated": False, "error": str(exc)}

    ascii_strings = _extract_ascii_strings(data, min_length)
    utf16_strings = _extract_utf16le_strings(data, min_length)

    combined = ascii_strings + utf16_strings
    combined = _deduplicate_preserve_order(combined)
    if limit:
        combined = combined[:limit]

    return {
        "strings": combined,
        "truncated": truncated,
        "bytes_processed": len(data),
        "file_size": file_size
    }


def _looks_like_base64(candidate: str, min_length: int = 32) -> bool:
    stripped = candidate.strip()
    if len(stripped) < min_length:
        return False
    if len(stripped) % 4 != 0:
        return False
    if stripped.count('=') > 2:
        return False
    return all(char in BASE64_ALPHABET for char in stripped)


def analyze_strings(strings: List[str]) -> Dict[str, Any]:
    """Анализирует извлеченные строки и ищет потенциальные IoC."""
    keyword_hits: Dict[str, List[str]] = collections.defaultdict(list)
    urls_found = set()
    ips_found = set()
    emails_found = set()
    base64_candidates = []

    total_length = 0
    max_length = 0

    for string in strings:
        lowered = string.lower()
        total_length += len(string)
        max_length = max(max_length, len(string))

        for keyword in SUSPICIOUS_KEYWORDS:
            if keyword in lowered and len(keyword_hits[keyword]) < 5:
                keyword_hits[keyword].append(string)

        for url in URL_REGEX_SCRIPT.findall(string):
            urls_found.add(url)
        for ip in IP_REGEX_SCRIPT.findall(string):
            ips_found.add(ip)
        for email in EMAIL_REGEX.findall(string):
            emails_found.add(email)

        if _looks_like_base64(string):
            if len(base64_candidates) < 10:
                base64_candidates.append(string)

    average_length = (total_length / len(strings)) if strings else 0

    sample_strings = strings[:STRING_SAMPLE_SIZE]

    return {
        "total_extracted": len(strings),
        "average_length": average_length,
        "longest_length": max_length,
        "keyword_hits": dict(keyword_hits),
        "network_indicators": {
            "urls": sorted(urls_found),
            "ips": sorted(ips_found)
        },
        "emails": sorted(emails_found),
        "base64_candidates": base64_candidates,
        "sample": sample_strings
    }


def _enhanced_pe_checks(base_analysis: Dict[str, Any]) -> Dict[str, Any]:
    """Дополнительные эвристики для PE-файлов."""
    if not isinstance(base_analysis, dict):
        return {}

    sections = base_analysis.get("sections") or []
    imports = base_analysis.get("imports") or []

    high_entropy_sections = []
    packer_sections = []
    entropies = []

    for section in sections:
        name = (section.get("name") or "").strip()
        entropy = section.get("entropy")
        if isinstance(entropy, (int, float)):
            entropies.append(entropy)
            if entropy >= HIGH_ENTROPY_THRESHOLD:
                high_entropy_sections.append({"name": name, "entropy": entropy})

        if name.lower() in PACKER_SECTION_NAMES_LOWER:
            packer_sections.append(name)

    suspicious_imports = []
    suspicious_dlls = []
    for imp in imports:
        dll = (imp.get("dll") or "").lower()
        func_name = (imp.get("name") or "")
        if func_name and func_name in SUSPICIOUS_IMPORT_NAMES:
            suspicious_imports.append({
                "dll": imp.get("dll"),
                "name": func_name,
                "address": imp.get("address")
            })
        if dll in SUSPICIOUS_DLL_NAMES:
            suspicious_dlls.append(imp.get("dll"))

    suspicious_imports = _deduplicate_preserve_order([
        f"{entry.get('dll')}!{entry.get('name')}" for entry in suspicious_imports
    ])
    suspicious_dlls = _deduplicate_preserve_order(suspicious_dlls)

    avg_entropy = sum(entropies) / len(entropies) if entropies else None
    low_import_count = 0 < len(imports) < 5

    possible_packer = bool(packer_sections or (len(high_entropy_sections) >= 2 and low_import_count))

    return {
        "high_entropy_sections": high_entropy_sections,
        "packer_sections": packer_sections,
        "suspicious_imports": suspicious_imports,
        "suspicious_dlls": suspicious_dlls,
        "average_entropy": avg_entropy,
        "total_sections": len(sections),
        "total_imports": len(imports),
        "low_import_count": low_import_count,
        "possible_packer": possible_packer
    }


def _compose_enhanced_summary(file_type: str,
                              base_analysis: Dict[str, Any],
                              string_report: Dict[str, Any]) -> Dict[str, Any]:
    """Формирует общий вывод расширенных проверок и итоговый скоринг."""
    score = 0
    reasons = []
    result: Dict[str, Any] = {
        "strings": string_report
    }

    pe_heuristics = {}
    if file_type == "pe":
        pe_heuristics = _enhanced_pe_checks(base_analysis)
        if pe_heuristics:
            result["pe_heuristics"] = pe_heuristics

        if pe_heuristics.get("possible_packer"):
            score += 3
            reasons.append("Обнаружены признаки упаковщика")
        if pe_heuristics.get("high_entropy_sections"):
            score += 2
            reasons.append("Высокая энтропия секций")
        if pe_heuristics.get("suspicious_imports"):
            score += min(3, len(pe_heuristics["suspicious_imports"]))
            reasons.append("Подозрительные импорты API")
        if pe_heuristics.get("low_import_count"):
            score += 1
            reasons.append("Низкое количество импортов")

    network_indicators = string_report.get("network_indicators", {})
    if network_indicators.get("urls") or network_indicators.get("ips"):
        score += 1
        reasons.append("Обнаружены сетевые индикаторы в строках")

    if string_report.get("keyword_hits"):
        score += 1
        reasons.append("Совпадения по ключевым словам угроз")

    if string_report.get("base64_candidates"):
        score += 1
        reasons.append("Подозрительные Base64-последовательности")

    severity = "low"
    if score >= 7:
        severity = "critical"
    elif score >= 4:
        severity = "high"
    elif score >= 2:
        severity = "medium"

    result["summary"] = {
        "score": score,
        "severity": severity,
        "reasons": reasons
    }

    return result

###############################################################################
# 2. EICAR signature + сигнатуры PE/ELF/PDF/ZIP/OLE/Скрипты
###############################################################################

EICAR_SIGNATURE = b"X5O!P%@AP[4\\PZX54(P^)7CC)7}$EICAR-STANDARD-ANTIVIRUS-TEST-FILE!$H+H*"

def is_eicar_file(filepath: str) -> bool:
    if not os.path.isfile(filepath):
        return False
    try:
        with open(filepath, "rb") as f:
            data = f.read()
            return (EICAR_SIGNATURE in data)
    except:
        return False

def _detect_by_signature(filepath: str) -> str:
    with open(filepath, "rb") as f:
        start = f.read(4)

    if start[:2] == b'MZ':
        return "pe"
    if start == b'\x7fELF':
        return "elf"
    if start == b'%PDF':
        return "pdf"
    if start[:2] == b'PK':
        return "zip"
    if start == b'\xd0\xcf\x11\xe0':
        return "ole"
    return "unknown"

def _inspect_ooxml_package(filepath: str) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "detected_type": "zip",
        "has_macros": False,
        "macro_streams": [],
        "embedded_objects": [],
        "external_relationships": [],
        "suspicious_relationships": [],
        "auto_open_indicators": [],
        "package_entries": 0,
    }
    try:
        with zipfile.ZipFile(filepath, "r") as archive:
            names = archive.namelist()
            lowered = {name.lower(): name for name in names}
            result["package_entries"] = len(names)

            macro_streams = [
                name for name in names
                if name.lower().endswith("vbaproject.bin")
            ]
            embedded_objects = [
                name for name in names
                if "/embeddings/" in name.lower()
            ]
            external_relationships: List[Dict[str, Any]] = []
            suspicious_relationships: List[Dict[str, Any]] = []
            auto_open_indicators: List[str] = []
            suspicious_rel_keywords = (
                "attachedtemplate",
                "external",
                "oleobject",
                "oleobjectlink",
                "hyperlink",
                "frame",
                "link",
            )
            auto_keywords = ("autoopen", "document_open", "autoexec", "attachedtemplate")

            for name in names:
                lower_name = name.lower()
                if lower_name.endswith(".rels"):
                    try:
                        root = ET.fromstring(archive.read(name))
                    except Exception:
                        continue
                    for rel in root.findall("{*}Relationship"):
                        rel_type = str(rel.attrib.get("Type", ""))
                        target = str(rel.attrib.get("Target", ""))
                        target_mode = str(rel.attrib.get("TargetMode", "Internal"))
                        item = {
                            "source": name,
                            "id": rel.attrib.get("Id"),
                            "type": rel_type,
                            "target": target,
                            "target_mode": target_mode,
                        }
                        if target_mode.lower() == "external" or target.lower().startswith(("http://", "https://", "\\\\")):
                            external_relationships.append(item)
                        combined = f"{rel_type} {target}".lower()
                        if any(keyword in combined for keyword in suspicious_rel_keywords):
                            suspicious_relationships.append(item)

                if lower_name.endswith((".xml", ".rels", ".vba")):
                    try:
                        xml_text = archive.read(name).decode("utf-8", errors="ignore").lower()
                    except Exception:
                        continue
                    for keyword in auto_keywords:
                        if keyword in xml_text and keyword not in auto_open_indicators:
                            auto_open_indicators.append(keyword)

            detected_type = "docm" if macro_streams else "docx" if "word/document.xml" in lowered else "zip"
            result.update(
                {
                    "detected_type": detected_type,
                    "has_macros": bool(macro_streams),
                    "macro_streams": macro_streams,
                    "embedded_objects": embedded_objects,
                    "external_relationships": external_relationships,
                    "suspicious_relationships": suspicious_relationships,
                    "auto_open_indicators": auto_open_indicators,
                }
            )
    except Exception as exc:
        result["error"] = str(exc)
    return result

def detect_file_type(filepath: str) -> str:
    # Сначала проверяем по расширению для скриптов
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".ps1":
        return "powershell_script"
    if ext == ".vbs":
        return "vbs_script"
    if ext == ".js":
        return "javascript_script"

    sigtype = _detect_by_signature(filepath)
    if sigtype == "pe":
        return "pe"
    elif sigtype == "elf":
        return "elf"
    elif sigtype == "pdf":
        return "pdf"
    elif sigtype == "ole":
        return "ole_doc"
    elif sigtype == "zip":
        ext = os.path.splitext(filepath)[1].lower()
        if ext in {".docx", ".docm"}:
            package_info = _inspect_ooxml_package(filepath)
            return package_info.get("detected_type", "docm" if ext == ".docm" else "docx")
        package_info = _inspect_ooxml_package(filepath)
        return package_info.get("detected_type", "zip")
    # Если сигнатуры нет, но расширение скриптовое (на всякий случай)
    if ext in [".ps1", ".vbs", ".js"]:
         logger.warning(f"Файл {filepath} без сигнатуры, но с расширением скрипта ({ext}). Пробуем анализ скрипта.")
         if ext == ".ps1": return "powershell_script"
         if ext == ".vbs": return "vbs_script"
         if ext == ".js": return "javascript_script"

    return "unknown"

###############################################################################
# 3. Дизассемблирование (Capstone)
###############################################################################

def _disassemble(data: bytes, arch='x86', is_64=True, start_addr=0) -> list:
    if not capstone:
        return ["capstone не установлен, дизассемблирование невозможно"]

    if arch == 'x86':
        if is_64:
            md = capstone.Cs(capstone.CS_ARCH_X86, capstone.CS_MODE_64)
        else:
            md = capstone.Cs(capstone.CS_ARCH_X86, capstone.CS_MODE_32)
    else:
        return ["Сейчас реализован только x86/x64"]

    md.detail = True
    instructions = []
    for i in md.disasm(data, start_addr):
        line = f"0x{i.address:x}:\t{i.mnemonic}\t{i.op_str}"
        instructions.append(line)
    return instructions

###############################################################################
# 4. Основные анализаторы: PE/ELF/PDF/DOCX/OLE
###############################################################################

def analyze_pe(filepath: str) -> dict:
    if not pefile:
        return {"error": "pefile не установлен"}
    try:
        pe = pefile.PE(filepath)
        info = {
            "entry_point": hex(pe.OPTIONAL_HEADER.AddressOfEntryPoint),
            "image_base": hex(pe.OPTIONAL_HEADER.ImageBase),
            "sections": [],
            "imports": [],
            "disassembly": []
        }
        for section in pe.sections:
            info["sections"].append({
                "name": section.Name.decode().strip(),
                "size": section.SizeOfRawData,
                "entropy": calculate_entropy(section.get_data())
            })
        if hasattr(pe, 'DIRECTORY_ENTRY_IMPORT'):
            for entry in pe.DIRECTORY_ENTRY_IMPORT:
                for imp in entry.imports:
                    info["imports"].append({
                        "dll": entry.dll.decode(),
                        "name": imp.name.decode() if imp.name else None,
                        "address": hex(imp.address)
                    })

        # --- Save Import Graph ---
        try:
            import matplotlib.pyplot as plt
        except ImportError:
            plt = None
        import os

        if nx and plt:
            G = nx.DiGraph()
            for imp in info["imports"]:
                G.add_edge(imp["dll"], imp["name"] or "Unknown")

            graph_path = os.path.splitext(filepath)[0] + "_import_graph.png"
            pos = nx.spring_layout(G)
            plt.figure(figsize=(12, 8))
            nx.draw(G, pos, with_labels=True, node_color="#00F3FF", edge_color="#FF00FF", font_color="#0D0D15")
            plt.title("Import Graph")
            plt.savefig(graph_path, dpi=150)
            plt.close()

        # --- Save Entropy Chart ---
        section_names = [s["name"] for s in info["sections"]]
        entropies = [s["entropy"] for s in info["sections"]]

        if plt:
            chart_path = os.path.splitext(filepath)[0] + "_entropy_chart.png"
            plt.figure(figsize=(10, 5))
            plt.bar(section_names, entropies, color="#00FF88")
            plt.xlabel("Section")
            plt.ylabel("Entropy")
            plt.title("Section Entropy")
            for i, v in enumerate(entropies):
                plt.text(i, v + 0.1, f"{v:.2f}", ha='center', fontsize=10)
            plt.savefig(chart_path, dpi=150)
            plt.close()

        return info
    except Exception as e:
        return {"error": f"Ошибка анализа PE: {e}"}

def get_import_graph(filepath: str):
    """
    Возвращает граф импортов PE-файла как networkx.DiGraph либо список рёбер,
    если networkx недоступен. Бросает исключение, если pefile отсутствует или файл не PE.
    """
    if not pefile:
        raise RuntimeError("pefile не установлен")
    pe = pefile.PE(filepath)
    imports = []
    if hasattr(pe, 'DIRECTORY_ENTRY_IMPORT'):
        for entry in pe.DIRECTORY_ENTRY_IMPORT:
            for imp in entry.imports:
                imports.append((entry.dll.decode(), (imp.name.decode() if imp.name else "Unknown")))
    if nx:
        G = nx.DiGraph()
        G.add_edges_from(imports)
        return G
    return imports

def get_entropy_graph(filepath: str):
    """
    Возвращает «граф» энтропии секций PE-файла в виде:
    - networkx.Graph c узлами Section и узлами Entropy:xx, соединёнными ребром; либо
    - списка рёбер [(section, f"entropy:{val:.2f}")].
    """
    if not pefile:
        raise RuntimeError("pefile не установлен")
    pe = pefile.PE(filepath)
    sections = []
    for section in pe.sections:
        name = section.Name.decode(errors='ignore').strip() or "<noname>"
        ent = calculate_entropy(section.get_data())
        sections.append((name, ent))
    edges = [(name, f"entropy:{ent:.2f}") for name, ent in sections]
    if nx:
        G = nx.Graph()
        G.add_edges_from(edges)
        return G
    return edges

def analyze_pdf(filepath: str) -> dict:
    if not PyPDF2:
        return {"error": "PyPDF2 не установлен"}
    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        num_pages = len(reader.pages)
        meta = {}
        if reader.metadata:
            for k, v in reader.metadata.items():
                meta[k] = str(v)
    return {
        "pages": num_pages,
        "metadata": meta
    }

def analyze_docx(filepath: str) -> dict:
    package_info = _inspect_ooxml_package(filepath)
    result = dict(package_info)
    preview: List[str] = []
    par_count = 0
    tbl_count = 0

    if docx:
        try:
            document = docx.Document(filepath)
            par_count = len(document.paragraphs)
            tbl_count = len(document.tables)
            for p in document.paragraphs[:8]:
                text = p.text.strip()
                if text:
                    preview.append(text)
        except Exception as exc:
            result["parser_warning"] = str(exc)
    else:
        result["parser_warning"] = "python-docx не установлен"

    if not preview:
        try:
            with zipfile.ZipFile(filepath, "r") as archive:
                xml_data = archive.read("word/document.xml").decode("utf-8", errors="ignore")
            fragments = re.findall(r">([^<>]{3,120})<", xml_data)
            preview = [fragment.strip() for fragment in fragments if fragment.strip()][:8]
        except Exception:
            preview = []

    result.update(
        {
            "paragraph_count": par_count,
            "table_count": tbl_count,
            "preview": preview,
        }
    )
    return result

def analyze_ole_doc(filepath: str) -> dict:
    if not VBA_Parser:
        return {"error": "oletools (VBA_Parser) не установлен"}
    res = {
        "has_macros": False,
        "macro_infos": [],
        "suspicious": [],
        "metadata": {}
    }
    parser = VBA_Parser(filepath)
    try:
        if parser.detect_vba_macros():
            res["has_macros"] = True
            for (filename, stream_path, vba_fn, vba_code) in parser.extract_macros():
                res["macro_infos"].append({
                    "stream_path": stream_path,
                    "vba_filename": vba_fn,
                    "code_length": len(vba_code)
                })
            for kw_type, keyword, descr in parser.analyze_macros():
                suspicious_str = f"{kw_type}: {keyword} => {descr}"
                res["suspicious"].append(suspicious_str)
        if parser.metadata:
            meta_d = {}
            for k, v in parser.metadata.items():
                meta_d[k] = str(v)
            res["metadata"] = meta_d
    finally:
        parser.close()
    return res

def analyze_elf(filepath: str) -> dict:
    if not ELFFile:
        return {"error": "pyelftools не установлен"}
    info = {
        "entry_point": None,
        "arch": None,
        "sections": [],
        "disassembly": []
    }
    with open(filepath, "rb") as f:
        elf = ELFFile(f)
        info["entry_point"] = hex(elf["e_entry"])
        info["arch"] = str(elf["e_machine"])
        for section in elf.iter_sections():
            info["sections"].append({
                "name": section.name,
                "size": section["sh_size"],
                "addr": hex(section["sh_addr"])
            })
        if capstone:
            text_sec = None
            for section in elf.iter_sections():
                if section.name == '.text':
                    text_sec = section
                    break
            if text_sec is not None:
                offset_in_section = elf["e_entry"] - text_sec["sh_addr"]
                if 0 <= offset_in_section < text_sec.data_size:
                    data = text_sec.data()
                    code_64 = data[offset_in_section : offset_in_section+64]
                    ds = _disassemble(code_64, arch='x86', is_64=True,
                                      start_addr=elf["e_entry"])
                    info["disassembly"] = ds
                else:
                    info["disassembly"].append(
                        "Entry point не попал в .text или за границами"
                    )
            else:
                info["disassembly"].append("Не найдена .text секция")
    return info

###############################################################################
# 4.2. Анализ скриптов (PS1, VBS, JS)
###############################################################################

# --- Паттерны для поиска ---
SCRIPT_PATTERNS = {
    "powershell": {
        "suspicious_cmdlets": [
            re.compile(r"Invoke-(Expression|Command|WMIMethod|RestMethod)", re.I),
            re.compile(r"(New-Object System\\.Net\\.WebClient)\\.Download(String|File)", re.I),
            re.compile(r"Start-Process", re.I),
            re.compile(r"Get-WMIObject", re.I),
            re.compile(r"Set-ItemProperty", re.I), # Часто для реестра/автозапуска
            re.compile(r"\bIEX\b", re.I),
        ],
        "obfuscation": [
            re.compile(r"-EncodedCommand\s+[A-Za-z0-9+/=]{20,}", re.I), # Base64 команда
            re.compile(r"FromBase64String", re.I),
            re.compile(r"\.(Replace|Join)\s*\(", re.I), # Конкатенация/замена
            re.compile(r"[GC]M\s+"  # Сокращения Get-Command/Measure-Command
                       r"(\*\s+)?"  # Возможно со звездочкой
                       r'([\"\'][^\"\']+)?', re.I), # Alias + string
            re.compile(r"\bxor\b", re.I),
        ]
    },
    "vbs": {
        "suspicious_objects": [
            re.compile(r"CreateObject\s*\(\s*['\"](WScript\.Shell|Scripting\.FileSystemObject|MSXML2\.XMLHTTP|Adodb\.Stream)['\"]\s*\)", re.I),
            re.compile(r"GetObject\s*\(\s*['\"]winmgmts:[\"']\s*\)", re.I), # WMI
            re.compile(r"Execute", re.I),
            re.compile(r"Eval", re.I),
        ],
        "obfuscation": [
            re.compile(r"Chr\(\d+\)", re.I), # Конструирование строк
            re.compile(r"&\s*['\"]", re.I), # Конкатенация
            re.compile(r"\b(Split|Join|Replace)\s*\(", re.I),
        ]
    },
    "javascript": {
        "suspicious_functions": [
            re.compile(r"eval\s*\(", re.I),
            re.compile(r"document\.write", re.I),
            re.compile(r"new\s+(ActiveXObject|XMLHttpRequest)", re.I),
            re.compile(r"(window\.)?location\.href"), # Редиректы
            re.compile(r"setTimeout|setInterval", re.I), # Задержки/циклы
            re.compile(r"unescape|escape", re.I), # Устаревшее, но бывает
        ],
        "obfuscation": [
            re.compile(r"String\.fromCharCode\s*\(", re.I),
            re.compile(r"\+\s*['\"]", re.I), # Конкатенация строк
            re.compile(r"parseInt\s*\(\"[0-9A-Fa-f]+\"\s*,\s*16\)", re.I), # Hex decode
            re.compile(r"atob\s*\(", re.I), # Base64 decode
            re.compile(r"split\s*\([\".\"]+\)\.join\s*\([\".\"]+\)", re.I), # Replace-like obfuscation
        ]
    }
}

def _analyze_script_content(content: str, script_type: str) -> Dict[str, List[str]]:
    """
    Общая функция для анализа содержимого скрипта на основе паттернов.
    :param content: Содержимое скрипта (строка).
    :param script_type: 'powershell', 'vbs' или 'javascript'.
    :return: Словарь с найденными признаками.
    """
    results = {
        "suspicious_constructs": [],
        "obfuscation_indicators": [],
        "network_indicators_found": [] # URL/IP
    }
    patterns = SCRIPT_PATTERNS.get(script_type, {})
    if not patterns:
        return results

    # Поиск подозрительных конструкций
    for category, regex_list in patterns.items():
        if category not in results: continue # Игнорируем, если нет такой категории в results
        key_name = "suspicious_constructs" if category in ["suspicious_cmdlets", "suspicious_objects", "suspicious_functions"] else "obfuscation_indicators"

        for regex in regex_list:
            # Ищем все совпадения, берем первые N строк для контекста
            matches = regex.finditer(content)
            for match in matches:
                 # Берем строку, где найдено совпадение
                 start_line = content.rfind('\n', 0, match.start()) + 1
                 end_line = content.find('\n', match.end())
                 if end_line == -1: end_line = len(content)
                 context_line = content[start_line:end_line].strip()
                 # Добавляем уникальные строки с контекстом
                 if context_line and context_line not in results[key_name]:
                     results[key_name].append(context_line)

    # Поиск сетевых индикаторов
    urls = URL_REGEX_SCRIPT.findall(content)
    ips = IP_REGEX_SCRIPT.findall(content)
    if urls:
        results["network_indicators_found"].extend(list(set(urls))) # Уникальные URL
    if ips:
        results["network_indicators_found"].extend(list(set(ips))) # Уникальные IP

    return results

def analyze_powershell(filepath: str) -> dict:
    """ Анализ PowerShell скрипта (.ps1). """
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return _analyze_script_content(content, "powershell")
    except Exception as e:
        return {"error": f"Ошибка чтения/анализа PS1: {e}"}

def analyze_vbs(filepath: str) -> dict:
    """ Анализ VBScript (.vbs). """
    try:
        # VBS часто в других кодировках, пробуем несколько
        encodings_to_try = ['utf-8', 'cp1251', 'cp1252', 'latin-1']
        content = None
        for enc in encodings_to_try:
            try:
                with open(filepath, "r", encoding=enc) as f:
                    content = f.read()
                break # Успешно прочитали
            except UnicodeDecodeError:
                continue
        if content is None:
             raise ValueError("Не удалось определить кодировку VBS файла")
        return _analyze_script_content(content, "vbs")
    except Exception as e:
        return {"error": f"Ошибка чтения/анализа VBS: {e}"}

def analyze_javascript(filepath: str) -> dict:
    """ Анализ JavaScript (.js). """
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return _analyze_script_content(content, "javascript")
    except Exception as e:
        return {"error": f"Ошибка чтения/анализа JS: {e}"}

###############################################################################
# 5. Yara Integration
###############################################################################

def _load_yara_rules_with_status(rules_dir: str) -> Tuple[Optional[object], Dict[str, Any]]:
    status: Dict[str, Any] = {
        "status": "disabled",
        "available": bool(yara),
        "rules_dir": os.path.abspath(rules_dir) if rules_dir else None,
        "rules_loaded": 0,
        "rule_files": [],
        "error": None,
    }
    if not yara:
        status["status"] = "engine_missing"
        return None, status
    if not rules_dir:
        status["status"] = "disabled"
        return None, status
    if not os.path.isdir(rules_dir):
        status["status"] = "directory_missing"
        logger.info(f"Директория с Yara-правилами не найдена: {rules_dir}")
        return None, status

    filepaths = {}
    try:
        for filename in os.listdir(rules_dir):
            if filename.lower().endswith(('.yar', '.yara')):
                filepath = os.path.join(rules_dir, filename)
                namespace = os.path.splitext(filename)[0]
                filepaths[namespace] = filepath

        status["rule_files"] = sorted(os.path.abspath(path) for path in filepaths.values())
        status["rules_loaded"] = len(filepaths)

        if not filepaths:
            status["status"] = "no_rules"
            logger.warning(f"Не найдено Yara-правил в {rules_dir}")
            return None, status

        logger.debug(f"Компиляция Yara-правил из: {filepaths}")
        rules = yara.compile(filepaths=filepaths)
        logger.info(f"Успешно скомпилировано {len(filepaths)} файлов Yara-правил.")
        status["status"] = "ready"
        return rules, status

    except yara.Error as e:
        status["status"] = "compile_error"
        status["error"] = str(e)
        logger.error(f"Ошибка компиляции Yara-правил: {e}")
        return None, status
    except Exception as e:
        status["status"] = "load_error"
        status["error"] = str(e)
        logger.error(f"Неожиданная ошибка при загрузке Yara-правил: {e}")
        return None, status


def _load_yara_rules(rules_dir: str) -> Optional[object]:
    """
    Загружает и компилирует Yara-правила из указанной директории.
    Имена файлов (без расширения) используются как пространства имен.
    :param rules_dir: Путь к директории с .yar/.yara файлами.
    :return: Скомпилированные правила yara.Rules или None в случае ошибки.
    """
    rules, _ = _load_yara_rules_with_status(rules_dir)
    return rules

def _scan_with_yara(filepath: str, compiled_rules) -> List[Dict[str, Any]]:
    """
    Сканирует файл с помощью скомпилированных Yara-правил.
    :param filepath: Путь к файлу для сканирования.
    :param compiled_rules: Скомпилированные правила.
    :return: Список словарей с совпавшими правилами [{'rule': '...', 'tags': [...], 'meta': {...}}].
    """
    if not compiled_rules:
        return []

    matches_list = []
    try:
        matches = compiled_rules.match(filepath)
        for match in matches:
            matches_list.append({
                'rule': match.rule,
                'namespace': match.namespace,
                'tags': match.tags,
                'meta': match.meta
            })
        if matches_list:
             logger.debug(f"Yara: найдено {len(matches_list)} совпадений для {filepath}")

    except yara.Error as e:
        logger.error(f"Ошибка при Yara-сканировании файла {filepath}: {e}")
    except Exception as e:
        logger.error(f"Неожиданная ошибка при Yara-сканировании {filepath}: {e}")

    return matches_list

###############################################################################
# 6. Основная точка входа: static_analysis(filepath)
###############################################################################

def static_analysis(filepath: str, yara_rules_dir: Optional[str] = './yara_rules') -> dict:
    """
    Универсальная функция статического анализа:
    1) Проверяем наличие файла
    2) Считаем MD5/SHA256
    3) Определяем формат (PE, ELF, PDF, docx, ole_doc, zip, ps1, vbs, js, unknown)
    4) Вызываем соответствующий анализ
    5) Если yara доступен и правила есть, сканируем файл.
    6) Извлекаем строки и запускаем расширенные проверки (эвристики, IoC).
    7) Возвращаем общий словарь
    :param filepath: Путь к файлу.
    :param yara_rules_dir: Путь к директории с Yara-правилами (опционально).
    """
    if not os.path.isfile(filepath):
        return {"error": f"File '{filepath}' не найден"}

    # 1) Хэши
    hashes = get_file_hashes(filepath)

    # 2) Тип файла
    ftype = detect_file_type(filepath)
    analysis_res = {}

    # 3) Вызов соответствующего анализа
    if ftype == "pe":
        analysis_res = analyze_pe(filepath)
    elif ftype == "elf":
        analysis_res = analyze_elf(filepath)
    elif ftype == "pdf":
        analysis_res = analyze_pdf(filepath)
    elif ftype in {"docx", "docm"}:
        analysis_res = analyze_docx(filepath)
    elif ftype == "ole_doc":
        analysis_res = analyze_ole_doc(filepath)
    elif ftype == "powershell_script":
        analysis_res["script_analysis"] = analyze_powershell(filepath)
    elif ftype == "vbs_script":
        analysis_res["script_analysis"] = analyze_vbs(filepath)
    elif ftype == "javascript_script":
        analysis_res["script_analysis"] = analyze_javascript(filepath)
    elif ftype == "zip":
        analysis_res["info"] = "ZIP-архив (не docx?). Статический анализ не реализован."
    else:
        # 4) Если unknown, проверяем EICAR
        if is_eicar_file(filepath):
            analysis_res["info"] = "EICAR Test File (тестовая сигнатура)."
        else:
            analysis_res["info"] = "Неизвестный формат или поддержка не реализована."

    # 5) Yara сканирование
    yara_matches = []
    if yara_rules_dir:
        compiled_rules, yara_status = _load_yara_rules_with_status(yara_rules_dir)
    else:
        compiled_rules, yara_status = None, {
            "status": "disabled",
            "available": bool(yara),
            "rules_dir": None,
            "rules_loaded": 0,
            "rule_files": [],
            "error": None,
        }
    if compiled_rules:
        yara_matches = _scan_with_yara(filepath, compiled_rules)
    yara_status["matches"] = len(yara_matches)

    # 6) Расширенные проверки (строки, эвристики)
    string_extraction = extract_strings(filepath)
    string_analysis = analyze_strings(string_extraction.get("strings", []))
    string_analysis["extraction_meta"] = {
        "truncated": string_extraction.get("truncated"),
        "bytes_processed": string_extraction.get("bytes_processed"),
        "file_size": string_extraction.get("file_size"),
        "errors": string_extraction.get("error")
    }

    enhanced_summary = _compose_enhanced_summary(ftype, analysis_res, string_analysis)

    # Итоговый словарь
    result = {
        "filepath": filepath,
        "hashes": hashes,
        "file_type": ftype,
        "analysis": analysis_res,
        "yara_matches": yara_matches, # Добавляем результаты Yara
        "yara_status": yara_status,
        "enhanced_checks": enhanced_summary
    }
    return result

###############################################################################
# 7. Новая функция: static_analysis_to_json
###############################################################################

def static_analysis_to_json(filepath: str, outpath: str) -> dict:
    """
    Выполняет статический анализ (static_analysis), затем
    сохраняет результат в JSON (outpath).
    Возвращает тот же словарь результата.
    """
    import json
    res = static_analysis(filepath)
    with open(outpath, "w", encoding="utf-8") as f:
        json.dump(res, f, indent=2, ensure_ascii=False)
    return res

###############################################################################
# 8. Локальный тест
###############################################################################

if __name__ == "__main__":
    import sys
    import json # Добавлено для вывода
    logging.basicConfig(level=logging.DEBUG)

    # Определяем директорию с правилами Yara рядом со скриптом
    script_dir = os.path.dirname(os.path.abspath(__file__))
    default_yara_dir = os.path.join(script_dir, 'yara_rules')

    if len(sys.argv) > 1:
        # Анализ указанного файла
        testf = sys.argv[1]
        yara_dir = default_yara_dir
        # Можно добавить опциональный аргумент для пути к правилам
        # if len(sys.argv) > 2: yara_dir = sys.argv[2]
        out = os.path.splitext(os.path.basename(testf))[0] + "_static_result.json"
        r = static_analysis(testf, yara_rules_dir=yara_dir) # Передаем путь к правилам
        print(f"Результат анализа сохранен в {out}")
        with open(out, "w", encoding="utf-8") as f:
             json.dump(r, f, indent=2, ensure_ascii=False)
        # print(json.dumps(r, indent=2, ensure_ascii=False)) # Вывод в консоль
    else:
        # Или просто тест из списка
        test_files = [
            "calc.exe",
            "test.elf",
            "document.pdf",
            "report.docx",
            "legacy.doc",
            "eicar.com",
            "test_script.ps1", # Пример скрипта
            "test_script.vbs", # Пример скрипта
            "test_script.js"   # Пример скрипта
        ]
        for tf in test_files:
            # Создадим пустые файлы для теста, если их нет
            if tf.endswith(('.ps1', '.vbs', '.js')) and not os.path.exists(tf):
                try:
                    with open(tf, 'w') as f_script:
                        if tf.endswith('.ps1'):
                            f_script.write('Write-Host "Test PS1"\n$ip = "192.168.1.1" # Test')
                        if tf.endswith('.vbs'):
                            f_script.write('"Test VBS\nWScript.Echo \"Hello\"\nSet objShell = CreateObject(\"WScript.Shell\")')
                        if tf.endswith('.js'):
                            f_script.write('// Test JS\nvar x = "Hello";\neval("console.log(\'test eval\')");')
                    logger.info(f"Создан пустой тестовый файл: {tf}")
                except Exception as e_create:
                    logger.error(f"Не удалось создать тестовый файл {tf}: {e_create}")

            if os.path.exists(tf):
                print(f"\n=== Статический анализ {tf} ===")
                info = static_analysis(tf, yara_rules_dir=default_yara_dir) # Передаем путь
                print(json.dumps(info, indent=2, ensure_ascii=False)) # Вывод JSON
            else:
                print(f"[Warning] Файл {tf} не найден, пропускаем.")



# LINE1
# LINE2
# LINE3
